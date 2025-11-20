"""DOCX document parser with structure preservation."""

import logging
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument

from app.core.exceptions import FileCorruptedError, ParsingError
from app.services.parsers.base import BaseDocumentParser, ParsedUnit

logger = logging.getLogger(__name__)


class DOCXParser(BaseDocumentParser):
    """
    DOCX document parser with structure preservation.

    Parsing strategy:
    - Use built-in paragraph structure from python-docx
    - Detect headings from style names (Heading 1, Heading 2, etc.)
    - Preserve heading hierarchy
    - Parse tables as separate units
    """

    def __init__(self) -> None:
        """Initialize DOCX parser."""
        self.logger = logger

    def validate_file(self, file_path: Path) -> bool:
        """
        Validate DOCX file is readable.

        Args:
            file_path: Path to DOCX file

        Returns:
            True if file is valid

        Raises:
            FileCorruptedError: If DOCX cannot be read
        """
        try:
            doc = DocxDocument(file_path)
            # Check if we can read paragraphs
            para_count = len(doc.paragraphs)
            self.logger.debug(f"DOCX validation successful: {para_count} paragraphs")
            return True
        except Exception as e:
            self.logger.error(f"DOCX validation failed for {file_path}: {e}")
            raise FileCorruptedError(f"Cannot read DOCX file: {e}")

    def parse(self, file_path: Path) -> list[ParsedUnit]:
        """
        Parse DOCX into atomic units.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of ParsedUnit objects

        Raises:
            ParsingError: If parsing fails
            FileCorruptedError: If file is corrupted
        """
        self.validate_file(file_path)

        units: list[ParsedUnit] = []
        sequence = 0
        current_section: Optional[str] = None

        try:
            doc = DocxDocument(file_path)

            # Parse paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect unit type from style
                unit_type, level = self._detect_unit_type_from_style(para.style.name)

                # Track current section for context
                if unit_type == "heading":
                    current_section = text

                units.append(
                    ParsedUnit(
                        unit_type=unit_type,
                        text=text,
                        sequence=sequence,
                        level=level,
                        page_number=None,  # DOCX doesn't have explicit pages
                        section_title=(
                            current_section if unit_type != "heading" else None
                        ),
                        parent_sequence=None,
                        metadata={"style": para.style.name},
                    )
                )
                sequence += 1

            # Parse tables
            for table_idx, table in enumerate(doc.tables):
                table_units = self._parse_table(table, sequence, current_section)
                units.extend(table_units)
                sequence += len(table_units)

            self.logger.info(f"Parsed {len(units)} units from DOCX: {file_path.name}")
            return units

        except Exception as e:
            self.logger.error(f"DOCX parsing failed for {file_path}: {e}")
            raise ParsingError(f"Failed to parse DOCX: {e}")

    def _detect_unit_type_from_style(self, style_name: str) -> tuple[str, int]:
        """
        Detect unit type and level from Word style name.

        Args:
            style_name: Word style name (e.g., 'Heading 1', 'Normal')

        Returns:
            Tuple of (unit_type, level)
        """
        style_lower = style_name.lower()

        if "heading 1" in style_lower:
            return ("heading", 0)
        elif "heading 2" in style_lower:
            return ("heading", 1)
        elif "heading" in style_lower:
            return ("heading", 2)
        elif "list" in style_lower:
            return ("list_item", 1)
        else:
            return ("paragraph", 1)

    def _parse_table(
        self, table, start_sequence: int, section_title: Optional[str]
    ) -> list[ParsedUnit]:
        """
        Parse a Word table into atomic units.

        Args:
            table: python-docx table object
            start_sequence: Starting sequence number for table rows
            section_title: Current section title for context

        Returns:
            List of ParsedUnit objects for table rows
        """
        units: list[ParsedUnit] = []

        for row_idx, row in enumerate(table.rows):
            cells_text = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells_text)

            if not row_text.strip():
                continue

            units.append(
                ParsedUnit(
                    unit_type="table_row",
                    text=row_text,
                    sequence=start_sequence + row_idx,
                    level=2,
                    page_number=None,
                    section_title=section_title,
                    parent_sequence=None,
                    metadata={"table_row": row_idx, "cells": cells_text},
                )
            )

        return units
